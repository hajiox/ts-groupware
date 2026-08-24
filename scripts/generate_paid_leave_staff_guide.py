from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "TSG有給管理システム_全スタッフ向け運用案内.docx"
PDF_OUTPUT = ROOT / "docs" / "TSG有給管理システム_全スタッフ向け運用案内.pdf"
SCREENSHOT_DIR = ROOT / "public" / "manual" / "screenshots"
HOME_SCREENSHOT = SCREENSHOT_DIR / "paid-leave-home-mobile.png"
REQUEST_SCREENSHOT = SCREENSHOT_DIR / "paid-leave-request-mobile.png"
BALANCE_SCREENSHOT = SCREENSHOT_DIR / "paid-leave-balance-mobile.png"

FONT = "Yu Gothic"
NAVY = "17324D"
BLUE = "1769AA"
TEAL = "137A70"
GOLD = "B07A13"
INK = "182433"
MUTED = "5C6977"
PALE_BLUE = "EAF3FA"
PALE_TEAL = "E9F6F2"
PALE_GOLD = "FFF5DB"
PALE_RED = "FCEBEC"
BORDER = "CBD7E3"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 11906
MARGIN_DXA = 907
CONTENT_DXA = 10080
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Mm(width / 1440 * 25.4)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_run_font(run, size=10.5, color=INK, bold=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def format_paragraph(paragraph, before=0, after=5, line=1.2):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_text(doc, text, *, size=10.5, color=INK, bold=False, before=0, after=5, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    format_paragraph(paragraph, before, after)
    run = paragraph.add_run(text)
    set_run_font(run, size, color, bold)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_bullet(doc, text, *, bold_prefix=None):
    paragraph = doc.add_paragraph(style="List Bullet")
    format_paragraph(paragraph, 0, 3, 1.18)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, 10.2, INK, True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest, 10.2, INK)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 10.2, INK)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    format_paragraph(paragraph, 0, 4, 1.18)
    run = paragraph.add_run(text)
    set_run_font(run, 10.2, INK)
    return paragraph


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph, 2, 6, 1.2)
    paragraph.paragraph_format.left_indent = Mm(2)
    paragraph.paragraph_format.right_indent = Mm(2)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "6")
        border.set(qn("w:color"), accent)
        borders.append(border)
    p_pr.append(borders)
    label_run = paragraph.add_run(f"{label}  ")
    set_run_font(label_run, 10.6, accent, True)
    body_run = paragraph.add_run(text)
    set_run_font(body_run, 10.6, INK)
    return paragraph


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    fonts.set(qn("w:eastAsia"), FONT)
    run_properties.extend([fonts, color, underline])
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, 8.5, MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend([field_begin, instruction, field_end])


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(16)
    section.right_margin = Mm(16)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    heading_tokens = {
        "Heading 1": (15.5, NAVY, 13, 6),
        "Heading 2": (12.5, BLUE, 10, 5),
        "Heading 3": (11.2, NAVY, 7, 3),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.2)
        style.paragraph_format.left_indent = Mm(7.5)
        style.paragraph_format.first_line_indent = Mm(-3.8)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.18

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(header_paragraph, 0, 0, 1)
    left = header_paragraph.add_run("TSG  |  有給管理システム")
    set_run_font(left, 8.5, NAVY, True)
    right = header_paragraph.add_run("    株式会社テクニカルスタッフ")
    set_run_font(right, 8.5, MUTED)
    p_pr = header_paragraph._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BORDER)
    border.append(bottom)
    p_pr.append(border)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    add_page_number(footer_paragraph)


def add_title_block(doc):
    add_text(doc, "全スタッフ向け運用案内", size=9.5, color=TEAL, bold=True, after=2)
    title = doc.add_paragraph()
    format_paragraph(title, 0, 4, 1.0)
    run = title.add_run("TSG有給管理システム")
    set_run_font(run, 24, NAVY, True)
    subtitle = doc.add_paragraph()
    format_paragraph(subtitle, 0, 12, 1.15)
    run = subtitle.add_run("2026年8月1日から運用を開始します")
    set_run_font(run, 14, BLUE, True)

    for row_data in (
        ("発行日", "2026年7月26日", "対象", "株式会社テクニカルスタッフ 全スタッフ"),
        ("運用開始", "2026年8月1日", "管理システム", "TSG"),
    ):
        paragraph = doc.add_paragraph()
        format_paragraph(paragraph, 0, 1, 1.15)
        paragraph.paragraph_format.left_indent = Mm(2)
        paragraph.paragraph_format.right_indent = Mm(2)
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F5F8FB")
        p_pr.append(shading)
        label_run = paragraph.add_run(f"{row_data[0]}  ")
        set_run_font(label_run, 9.2, MUTED, True)
        value_run = paragraph.add_run(f"{row_data[1]}      ")
        set_run_font(value_run, 9.6, INK)
        second_label = paragraph.add_run(f"{row_data[2]}  ")
        set_run_font(second_label, 9.2, MUTED, True)
        second_value = paragraph.add_run(row_data[3])
        set_run_font(second_value, 9.6, INK)

    add_text(
        doc,
        "年次有給休暇の残日数、付与予定、申請、取得履歴をTSGで管理します。"
        "各自が自分の情報を確認でき、予定している有給はシフト希望回収と一緒に提出します。",
        size=11,
        after=9,
    )
    add_callout(
        doc,
        "重要",
        "予定して取得する有給はシフト希望回収から提出します。"
        "シフト確定後や突発休は、先に所属管理者へ連絡してください。"
        "口頭・電話・DMで連絡しただけでは、TSG上の申請は完了しません。",
        fill=PALE_GOLD,
        accent=GOLD,
    )


def add_summary_table(doc):
    add_heading(doc, "まず覚えてほしいこと", 1)
    rows = [
        ("予定して休む", "シフト希望回収から提出", "全休1日 / 半休0.5日"),
        ("締切前の変更", "同じ希望回収画面で修正", "保存し直せば更新"),
        ("突発休・確定後", "所属管理者へ先に連絡", "管理者確認後に反映"),
        ("打刻がない勤務日", "ホームの確認通知から回答", "実際の理由を選択"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2300, 3880, 3900])
    headers = ("場面", "行うこと", "システム上の扱い")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, 9.5, WHITE, True)
    mark_table_header(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            set_cell_border(cells[index])
            if index == 0:
                set_cell_shading(cells[index], PALE_BLUE)
            paragraph = cells[index].paragraphs[0]
            format_paragraph(paragraph, 0, 0, 1.12)
            run = paragraph.add_run(value)
            set_run_font(run, 9.4, INK, index == 0)
    set_table_geometry(table, [2300, 3880, 3900])


def add_visible_information(doc):
    add_heading(doc, "自分で確認できる情報", 1)
    items = [
        "現在の有給残日数、次回付与日、付与見込日数",
        "有効な付与分・期限、移行調整休暇の付与予定",
        "有給の申請・取得履歴",
        "月・年・入社後の欠勤回数と出勤率",
        "直近3か月の平均勤務時間と参考賃金",
    ]
    for item in items:
        add_bullet(doc, item)
    add_text(
        doc,
        "各スタッフが確認できるのは自分の情報だけです。"
        "管理上必要な権限を持つ担当者は、確認・承認・修正を行います。",
        size=9.6,
        color=MUTED,
        after=5,
    )


def add_planned_workflow(doc):
    add_heading(doc, "予定して有給を取得する場合", 1)
    steps = [
        "TSGホーム上部の「シフト希望回収」を開きます。",
        "「休み・有給希望」で「有給（全休）」または「有給（半休）」を選びます。",
        "カレンダーから希望日を選びます。",
        "選択内容と日数を確認し、提出期限までに「保存」を押します。",
        "提出期限前でシフトが確定していなければ、同じ画面から修正できます。",
    ]
    for step in steps:
        add_numbered(doc, step)
    add_callout(
        doc,
        "半休の日",
        "確定したシフトで勤務する時間に、通常どおり出勤・退勤を打刻してください。",
        fill=PALE_TEAL,
        accent=TEAL,
    )


def add_emergency_workflow(doc):
    add_heading(doc, "突発休・シフト確定後の変更", 1)
    steps = [
        "まず所属管理者へ連絡します。",
        "シフト確定後に有給へ変更する場合は、管理者がTSGへ登録します。",
        "確定シフトの勤務日に打刻がない場合、TSGホームに「勤怠の確認が必要です」と表示されます。",
        "通知を開き、「有給（全休）」「有給（半休）」「欠勤」「打刻忘れ」「勤務変更」から実際の状況を選び、必要に応じてメモを入力します。",
        "管理者が内容を確認・確定すると、有給または欠勤として正式に反映されます。",
    ]
    for step in steps:
        add_numbered(doc, step)
    add_callout(
        doc,
        "注意",
        "打刻忘れを有給・欠勤として回答しないでください。"
        "実際の状況と異なる回答をした場合は、すぐに所属管理者へ連絡してください。",
        fill=PALE_RED,
        accent="B33A42",
    )


def add_balance_access(doc):
    add_heading(doc, "有給残日数の確認", 1)
    add_text(doc, "TSGへログインした状態で「有給・欠勤」画面を開きます。", after=3)
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph, 0, 6, 1.1)
    add_hyperlink(paragraph, "https://v0-line-blush.vercel.app/leave", "https://v0-line-blush.vercel.app/leave")
    add_text(
        doc,
        "残日数が不足している場合、有給申請は登録できません。"
        "表示内容に疑問がある場合は、申請前に所属管理者へ確認してください。",
        size=9.7,
        color=MUTED,
        after=4,
    )


def add_page_break(doc):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_grant_rules(doc):
    add_heading(doc, "付与と残数のルール", 1)
    items = [
        "2026年7月31日までに確認した有給残数は、減らさず新システムへ引き継ぎます。",
        "2026年8月1日から、付与基準を全社一律の8月更新から、各自の入社日を基準とする方式へ順次移行します。",
        "移行により不利にならないよう、対象者には最初の個人付与日まで、月初に1日の「移行調整休暇」を付与します。",
        "移行調整休暇は法定有給とは別の会社休暇として管理し、付与予定日は本人画面に表示します。",
        "通常の有給は、入社日、勤続年数、勤務形態、所定労働日数に基づいて付与します。",
        "出勤率要件は、確定シフトと打刻情報から判定します。制度開始時の初回判定のみ、過去データが十分でないため要件を満たしたものとして扱います。",
        "有給を使用した場合は、期限の早い付与分から順に消化します。",
        "半休は0.5日として残数から差し引きます。",
    ]
    for item in items:
        add_bullet(doc, item)


def add_wage_rules(doc):
    add_heading(doc, "賃金の扱い", 1)
    items = [
        "月給制のスタッフは、給与計算上の登録内容に従います。",
        "パート・フルタイムパートは、直近3か月の平均勤務時間と登録時給を基準に参考賃金を算出します。",
        "システムに十分な勤務実績が蓄積するまでは、過去のシフト資料から登録した平均勤務時間を使用します。",
        "本人画面の金額は参考表示です。最終的な支給額は給与計算結果で確認してください。",
    ]
    for item in items:
        add_bullet(doc, item)


def add_faq(doc):
    add_heading(doc, "よくある質問", 1)
    faqs = [
        ("普通の休み希望と有給希望は何が違いますか",
         "普通の休み希望は有給残数を使用しません。"
         "「有給（全休）」または「有給（半休）」を選んだ日だけ、有給申請として処理されます。"),
        ("提出した後に変更できますか",
         "提出期限前かつシフト確定前であれば修正できます。期限後またはシフト確定後は、所属管理者へ連絡してください。"),
        ("有給を選んだ時点で残日数は減りますか",
         "選択しただけでは確定しません。シフト確定または管理者確認後に正式反映されます。"),
        ("スマートフォンを変えた場合、残数は消えますか",
         "消えません。有給情報は端末ではなくTSGアカウントごとに保存されます。"),
        ("表示されている残数や入社日が違います",
         "申請前に所属管理者へ連絡してください。確認後、人事・有給管理データを修正します。"),
    ]
    for question, answer in faqs:
        paragraph = doc.add_paragraph()
        format_paragraph(paragraph, 4, 1, 1.15)
        q_run = paragraph.add_run(f"Q  {question}")
        set_run_font(q_run, 10.1, NAVY, True)
        answer_paragraph = doc.add_paragraph()
        format_paragraph(answer_paragraph, 0, 4, 1.15)
        answer_paragraph.paragraph_format.left_indent = Mm(5)
        a_run = answer_paragraph.add_run(f"A  {answer}")
        set_run_font(a_run, 9.8, INK)


def add_screen_guide(doc):
    add_heading(doc, "「有給・欠勤」画面の見方", 1)
    rows = [
        ("有給残日数", "今日の時点で使用できる残日数。将来の付与予定は含みません。"),
        ("次回付与予定", "入社日基準で次に付与される予定日と見込日数。"),
        ("有効な有給", "付与日、期限、付与日数、残日数を付与単位で表示。"),
        ("移行調整休暇の予定", "付与日になると加算される会社独自休暇。予定表示中は現在残数に含みません。"),
        ("現在の出勤率", "確定シフトと打刻情報を基にした付与判定用の参考値。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    headers = ("表示", "意味")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run_font(run, 9.4, WHITE, True)
    mark_table_header(table.rows[0])
    for label, detail in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], PALE_BLUE)
        for index, value in enumerate((label, detail)):
            set_cell_border(cells[index])
            paragraph = cells[index].paragraphs[0]
            format_paragraph(paragraph, 0, 0, 1.12)
            run = paragraph.add_run(value)
            set_run_font(run, 9.2, INK, index == 0)
    set_table_geometry(table, [2700, 7380])


def add_contact(doc):
    add_heading(doc, "問い合わせ", 1)
    add_callout(
        doc,
        "連絡先",
        "操作方法、残日数、入社日、勤務形態、申請内容に疑問がある場合は、所属管理者へ連絡してください。"
        "緊急の欠勤連絡は、TSGへの入力だけで済ませず、これまでどおり所属管理者へ直接連絡してください。",
        fill=PALE_BLUE,
        accent=BLUE,
    )
    add_text(doc, "株式会社テクニカルスタッフ", size=10.5, color=NAVY, bold=True, after=7)
    source = doc.add_paragraph()
    format_paragraph(source, 4, 0, 1.1)
    source_run = source.add_run("参考: 厚生労働省「年次有給休暇」  ")
    set_run_font(source_run, 8.3, MUTED)
    add_hyperlink(
        source,
        "https://www.check-roudou.mhlw.go.jp/study/roudousya_yukyu.html",
        "https://www.check-roudou.mhlw.go.jp/study/roudousya_yukyu.html",
    )


def build():
    doc = Document()
    style_document(doc)
    add_title_block(doc)
    add_summary_table(doc)
    add_visible_information(doc)
    add_planned_workflow(doc)
    add_emergency_workflow(doc)
    add_balance_access(doc)
    add_page_break(doc)
    add_grant_rules(doc)
    add_wage_rules(doc)
    add_screen_guide(doc)
    add_faq(doc)
    add_contact(doc)

    properties = doc.core_properties
    properties.title = "TSG有給管理システム 全スタッフ向け運用案内"
    properties.subject = "2026年8月1日運用開始"
    properties.author = "株式会社テクニカルスタッフ"
    properties.keywords = "TSG, 有給, 年次有給休暇, 運用案内"
    properties.comments = "全スタッフ向け配布文書"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


def pdf_color(value):
    return colors.HexColor(f"#{value}")


def register_pdf_fonts():
    pdfmetrics.registerFont(TTFont("BIZUDGothic", r"C:\Windows\Fonts\BIZ-UDGothicR.ttc"))
    pdfmetrics.registerFont(TTFont("BIZUDGothic-Bold", r"C:\Windows\Fonts\BIZ-UDGothicB.ttc"))


def pdf_styles():
    base = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle(
            "BodyJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic",
            fontSize=9.6,
            leading=14,
            textColor=pdf_color(INK),
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "small": ParagraphStyle(
            "SmallJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic",
            fontSize=8.2,
            leading=11,
            textColor=pdf_color(MUTED),
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "kicker": ParagraphStyle(
            "KickerJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic-Bold",
            fontSize=8.8,
            leading=11,
            textColor=pdf_color(TEAL),
            spaceAfter=2,
            wordWrap="CJK",
        ),
        "title": ParagraphStyle(
            "TitleJP",
            parent=base["Title"],
            fontName="BIZUDGothic-Bold",
            fontSize=23,
            leading=27,
            textColor=pdf_color(NAVY),
            alignment=TA_LEFT,
            spaceAfter=3,
            wordWrap="CJK",
        ),
        "subtitle": ParagraphStyle(
            "SubtitleJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic-Bold",
            fontSize=13,
            leading=17,
            textColor=pdf_color(BLUE),
            spaceAfter=10,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "H1JP",
            parent=base["Heading1"],
            fontName="BIZUDGothic-Bold",
            fontSize=14,
            leading=18,
            textColor=pdf_color(NAVY),
            spaceBefore=9,
            spaceAfter=5,
            keepWithNext=True,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "H2JP",
            parent=base["Heading2"],
            fontName="BIZUDGothic-Bold",
            fontSize=11.2,
            leading=14,
            textColor=pdf_color(BLUE),
            spaceBefore=6,
            spaceAfter=3,
            keepWithNext=True,
            wordWrap="CJK",
        ),
        "question": ParagraphStyle(
            "QuestionJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic-Bold",
            fontSize=9.4,
            leading=13,
            textColor=pdf_color(NAVY),
            spaceBefore=4,
            spaceAfter=1,
            wordWrap="CJK",
        ),
        "answer": ParagraphStyle(
            "AnswerJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic",
            fontSize=9.2,
            leading=13,
            leftIndent=5 * mm,
            textColor=pdf_color(INK),
            spaceAfter=4,
            wordWrap="CJK",
        ),
        "table": ParagraphStyle(
            "TableJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic",
            fontSize=8.4,
            leading=11,
            textColor=pdf_color(INK),
            wordWrap="CJK",
        ),
        "table_bold": ParagraphStyle(
            "TableBoldJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic-Bold",
            fontSize=8.4,
            leading=11,
            textColor=pdf_color(INK),
            wordWrap="CJK",
        ),
        "table_header": ParagraphStyle(
            "TableHeaderJP",
            parent=base["BodyText"],
            fontName="BIZUDGothic-Bold",
            fontSize=8.5,
            leading=11,
            textColor=colors.white,
            alignment=TA_CENTER,
            wordWrap="CJK",
        ),
    }
    return styles


def pdf_header_footer(canvas, document):
    canvas.saveState()
    width, height = A4
    canvas.setStrokeColor(pdf_color(BORDER))
    canvas.setLineWidth(0.5)
    canvas.line(16 * mm, height - 12 * mm, width - 16 * mm, height - 12 * mm)
    canvas.setFont("BIZUDGothic-Bold", 7.7)
    canvas.setFillColor(pdf_color(NAVY))
    canvas.drawString(16 * mm, height - 9.5 * mm, "TSG  |  有給管理システム")
    canvas.setFont("BIZUDGothic", 7.7)
    canvas.setFillColor(pdf_color(MUTED))
    canvas.drawRightString(width - 16 * mm, height - 9.5 * mm, "株式会社テクニカルスタッフ")
    canvas.setFont("BIZUDGothic", 7.5)
    canvas.drawRightString(width - 16 * mm, 8 * mm, f"Page {document.page}")
    canvas.restoreState()


def pdf_callout(styles, label, text, fill=PALE_BLUE, accent=BLUE):
    content = Paragraph(
        f'<font name="BIZUDGothic-Bold" color="#{accent}">{label}</font>'
        f'　<font name="BIZUDGothic">{text}</font>',
        styles["body"],
    )
    table = Table([[content]], colWidths=[178 * mm], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), pdf_color(fill)),
        ("BOX", (0, 0), (-1, -1), 0.8, pdf_color(accent)),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    return table


def pdf_bullets(styles, items):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["body"]), leftIndent=3 * mm) for item in items],
        bulletType="bullet",
        start="circle",
        leftIndent=6 * mm,
        bulletFontName="BIZUDGothic",
        bulletFontSize=7,
        spaceAfter=3,
    )


def pdf_steps(styles, items):
    return ListFlowable(
        [ListItem(Paragraph(item, styles["body"]), leftIndent=3 * mm) for item in items],
        bulletType="1",
        leftIndent=7 * mm,
        bulletFontName="BIZUDGothic-Bold",
        bulletFontSize=8.5,
        bulletColor=pdf_color(BLUE),
        spaceAfter=3,
    )


def build_pdf():
    register_pdf_fonts()
    styles = pdf_styles()
    document = BaseDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=17 * mm,
        bottomMargin=14 * mm,
        title="TSG有給管理システム 全スタッフ向け運用案内",
        author="株式会社テクニカルスタッフ",
        subject="2026年8月1日運用開始",
    )
    frame = Frame(
        document.leftMargin,
        document.bottomMargin,
        document.width,
        document.height,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    document.addPageTemplates([
        PageTemplate(id="staff-guide", frames=[frame], onPage=pdf_header_footer),
    ])

    story = []
    story.extend([
        Spacer(1, 5 * mm),
        Paragraph("全スタッフ向け運用案内", styles["kicker"]),
        Paragraph("TSG有給管理システム", styles["title"]),
        Paragraph("2026年8月1日から運用を開始します", styles["subtitle"]),
    ])

    metadata = [
        [Paragraph("<b>発行日</b>　2026年7月26日", styles["table"]),
         Paragraph("<b>対象</b>　株式会社テクニカルスタッフ 全スタッフ", styles["table"])],
        [Paragraph("<b>運用開始</b>　2026年8月1日", styles["table"]),
         Paragraph("<b>管理システム</b>　TSG", styles["table"])],
    ]
    meta_table = Table(metadata, colWidths=[89 * mm, 89 * mm], hAlign="LEFT")
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), pdf_color("F5F8FB")),
        ("GRID", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([
        meta_table,
        Spacer(1, 4 * mm),
        Paragraph(
            "年次有給休暇の残日数、付与予定、申請、取得履歴をTSGで管理します。"
            "各自が自分の情報を確認でき、予定している有給はシフト希望回収と一緒に提出します。",
            styles["body"],
        ),
        Spacer(1, 1.5 * mm),
        pdf_callout(
            styles,
            "重要",
            "予定して取得する有給はシフト希望回収から提出します。"
            "シフト確定後や突発休は、先に所属管理者へ連絡してください。"
            "口頭・電話・DMで連絡しただけでは、TSG上の申請は完了しません。",
            PALE_GOLD,
            GOLD,
        ),
        Paragraph("まず覚えてほしいこと", styles["h1"]),
    ])

    summary_rows = [
        [Paragraph(text, styles["table_header"]) for text in ("場面", "行うこと", "システム上の扱い")],
        [Paragraph("予定して休む", styles["table_bold"]), Paragraph("シフト希望回収から提出", styles["table"]), Paragraph("全休1日 / 半休0.5日", styles["table"])],
        [Paragraph("締切前の変更", styles["table_bold"]), Paragraph("同じ希望回収画面で修正", styles["table"]), Paragraph("保存し直せば更新", styles["table"])],
        [Paragraph("突発休・確定後", styles["table_bold"]), Paragraph("所属管理者へ先に連絡", styles["table"]), Paragraph("管理者確認後に反映", styles["table"])],
        [Paragraph("打刻がない勤務日", styles["table_bold"]), Paragraph("ホームの確認通知から回答", styles["table"]), Paragraph("実際の理由を選択", styles["table"])],
    ]
    summary = Table(summary_rows, colWidths=[41 * mm, 67 * mm, 70 * mm], repeatRows=1, hAlign="LEFT")
    summary.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), pdf_color(NAVY)),
        ("BACKGROUND", (0, 1), (0, -1), pdf_color(PALE_BLUE)),
        ("GRID", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([
        summary,
        Paragraph("自分で確認できる情報", styles["h1"]),
        pdf_bullets(styles, [
            "現在の有給残日数、次回付与日、付与見込日数",
            "有効な付与分・期限、移行調整休暇の付与予定",
            "有給の申請・取得履歴",
            "月・年・入社後の欠勤回数と出勤率",
            "直近3か月の平均勤務時間と参考賃金",
        ]),
        Paragraph(
            "各スタッフが確認できるのは自分の情報だけです。"
            "管理上必要な権限を持つ担当者は、確認・承認・修正を行います。",
            styles["small"],
        ),
        Paragraph("予定して有給を取得する場合", styles["h1"]),
        pdf_steps(styles, [
            "TSGホーム上部の「シフト希望回収」を開きます。",
            "「休み・有給希望」で「有給（全休）」または「有給（半休）」を選びます。",
            "カレンダーから希望日を選びます。",
            "選択内容と日数を確認し、提出期限までに「保存」を押します。",
            "提出期限前でシフトが確定していなければ、同じ画面から修正できます。",
        ]),
        pdf_callout(
            styles,
            "半休の日",
            "確定したシフトで勤務する時間に、通常どおり出勤・退勤を打刻してください。",
            PALE_TEAL,
            TEAL,
        ),
        PageBreak(),
        Paragraph("突発休・シフト確定後の変更", styles["h1"]),
        pdf_steps(styles, [
            "まず所属管理者へ連絡します。",
            "シフト確定後に有給へ変更する場合は、管理者がTSGへ登録します。",
            "確定シフトの勤務日に打刻がない場合、TSGホームに「勤怠の確認が必要です」と表示されます。",
            "通知を開き、「有給（全休）」「有給（半休）」「欠勤」「打刻忘れ」「勤務変更」から実際の状況を選び、必要に応じてメモを入力します。",
            "管理者が内容を確認・確定すると、有給または欠勤として正式に反映されます。",
        ]),
        pdf_callout(
            styles,
            "注意",
            "打刻忘れを有給・欠勤として回答しないでください。"
            "実際の状況と異なる回答をした場合は、すぐに所属管理者へ連絡してください。",
            PALE_RED,
            "B33A42",
        ),
        Paragraph("有給残日数の確認", styles["h1"]),
        Paragraph(
            'TSGへログインした状態で「有給・欠勤」画面を開きます。<br/>'
            '<link href="https://v0-line-blush.vercel.app/leave" color="#1769AA">'
            'https://v0-line-blush.vercel.app/leave</link>',
            styles["body"],
        ),
        Paragraph(
            "残日数が不足している場合、有給申請は登録できません。"
            "表示内容に疑問がある場合は、申請前に所属管理者へ確認してください。",
            styles["small"],
        ),
        Paragraph("付与と残数のルール", styles["h1"]),
        pdf_bullets(styles, [
            "2026年7月31日までに確認した有給残数は、減らさず新システムへ引き継ぎます。",
            "2026年8月1日から、付与基準を全社一律の8月更新から、各自の入社日を基準とする方式へ順次移行します。",
            "移行により不利にならないよう、対象者には最初の個人付与日まで、月初に1日の「移行調整休暇」を付与します。",
            "移行調整休暇は法定有給とは別の会社休暇として管理し、付与予定日は本人画面に表示します。",
            "通常の有給は、入社日、勤続年数、勤務形態、所定労働日数に基づいて付与します。",
            "出勤率要件は、確定シフトと打刻情報から判定します。制度開始時の初回判定のみ、過去データが十分でないため要件を満たしたものとして扱います。",
            "有給を使用した場合は、期限の早い付与分から順に消化します。",
            "半休は0.5日として残数から差し引きます。",
        ]),
        Paragraph("賃金の扱い", styles["h1"]),
        pdf_bullets(styles, [
            "月給制のスタッフは、給与計算上の登録内容に従います。",
            "パート・フルタイムパートは、直近3か月の平均勤務時間と登録時給を基準に参考賃金を算出します。",
            "システムに十分な勤務実績が蓄積するまでは、過去のシフト資料から登録した平均勤務時間を使用します。",
            "本人画面の金額は参考表示です。最終的な支給額は給与計算結果で確認してください。",
        ]),
        PageBreak(),
        Paragraph("「有給・欠勤」画面の見方", styles["h1"]),
    ])

    screen_rows = [
        [Paragraph(text, styles["table_header"]) for text in ("表示", "意味")],
        [Paragraph("有給残日数", styles["table_bold"]), Paragraph("今日の時点で使用できる残日数。将来の付与予定は含みません。", styles["table"])],
        [Paragraph("次回付与予定", styles["table_bold"]), Paragraph("入社日基準で次に付与される予定日と見込日数。", styles["table"])],
        [Paragraph("有効な有給", styles["table_bold"]), Paragraph("付与日、期限、付与日数、残日数を付与単位で表示。", styles["table"])],
        [Paragraph("移行調整休暇の予定", styles["table_bold"]), Paragraph("付与日になると加算される会社独自休暇。予定表示中は現在残数に含みません。", styles["table"])],
        [Paragraph("現在の出勤率", styles["table_bold"]), Paragraph("確定シフトと打刻情報を基にした付与判定用の参考値。", styles["table"])],
    ]
    screen_table = Table(screen_rows, colWidths=[45 * mm, 133 * mm], repeatRows=1, hAlign="LEFT")
    screen_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), pdf_color(NAVY)),
        ("BACKGROUND", (0, 1), (0, -1), pdf_color(PALE_BLUE)),
        ("GRID", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.extend([
        screen_table,
        Paragraph("よくある質問", styles["h1"]),
    ])

    faqs = [
        ("普通の休み希望と有給希望は何が違いますか",
         "普通の休み希望は有給残数を使用しません。"
         "「有給（全休）」または「有給（半休）」を選んだ日だけ、有給申請として処理されます。"),
        ("提出した後に変更できますか",
         "提出期限前かつシフト確定前であれば修正できます。期限後またはシフト確定後は、所属管理者へ連絡してください。"),
        ("有給を選んだ時点で残日数は減りますか",
         "選択しただけでは確定しません。シフト確定または管理者確認後に正式反映されます。"),
        ("スマートフォンを変えた場合、残数は消えますか",
         "消えません。有給情報は端末ではなくTSGアカウントごとに保存されます。"),
        ("表示されている残数や入社日が違います",
         "申請前に所属管理者へ連絡してください。確認後、人事・有給管理データを修正します。"),
    ]
    for question, answer in faqs:
        story.extend([
            KeepTogether([
                Paragraph(f"Q  {question}", styles["question"]),
                Paragraph(f"A  {answer}", styles["answer"]),
            ]),
        ])

    story.extend([
        Paragraph("問い合わせ", styles["h1"]),
        pdf_callout(
            styles,
            "連絡先",
            "操作方法、残日数、入社日、勤務形態、申請内容に疑問がある場合は、所属管理者へ連絡してください。"
            "緊急の欠勤連絡は、TSGへの入力だけで済ませず、これまでどおり所属管理者へ直接連絡してください。",
            PALE_BLUE,
            BLUE,
        ),
        Spacer(1, 3 * mm),
        Paragraph('<font name="BIZUDGothic-Bold" color="#17324D">株式会社テクニカルスタッフ</font>', styles["body"]),
        Spacer(1, 2 * mm),
        HRFlowable(width="100%", thickness=0.4, color=pdf_color(BORDER), spaceBefore=3, spaceAfter=4),
        Paragraph(
            '参考: 厚生労働省「年次有給休暇」<br/>'
            '<link href="https://www.check-roudou.mhlw.go.jp/study/roudousya_yukyu.html" color="#1769AA">'
            'https://www.check-roudou.mhlw.go.jp/study/roudousya_yukyu.html</link>',
            styles["small"],
        ),
    ])

    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.build(story)
    print(PDF_OUTPUT)


def add_docx_picture(cell, image_path: Path, width_mm: float):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))


def add_docx_cell_heading(cell, text: str):
    paragraph = cell.add_paragraph()
    format_paragraph(paragraph, 2, 5, 1.1)
    run = paragraph.add_run(text)
    set_run_font(run, 13, NAVY, True)


def add_docx_cell_bullet(cell, text: str):
    paragraph = cell.add_paragraph()
    format_paragraph(paragraph, 0, 3, 1.22)
    paragraph.paragraph_format.left_indent = Mm(4)
    paragraph.paragraph_format.first_line_indent = Mm(-3)
    run = paragraph.add_run(f"● {text}")
    set_run_font(run, 9.6, INK)


def set_paragraph_border_and_shading(paragraph, fill: str, accent: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "start", "bottom", "end"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), accent)
        borders.append(border)
    p_pr.append(borders)


def build_compact():
    doc = Document()
    style_document(doc)
    add_title_block_compact(doc)

    intro_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(intro_table, [5900, 4180])
    left, right = intro_table.rows[0].cells
    for cell in (left, right):
        set_cell_border(cell)
        set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    set_cell_shading(left, "F8FAFC")
    set_cell_shading(right, PALE_BLUE)

    add_docx_cell_heading(left, "8月1日から変わること")
    for item in (
        "残日数・付与予定・取得履歴をTSGで確認できます。",
        "予定して取る有給は、シフト希望回収と一緒に提出します。",
        "有給は全休1日、半休0.5日の2種類です。",
        "8月1日の付与分は、当日になると残日数へ加算されます。",
    ):
        add_docx_cell_bullet(left, item)

    add_docx_cell_heading(left, "有給を取るとき")
    for number, item in enumerate((
        "ホーム上部の「シフト希望回収」を開く",
        "「有給（全休）」または「有給（半休）」を選ぶ",
        "希望日を押す",
        "「保存」を押して提出する",
    ), 1):
        paragraph = left.add_paragraph()
        format_paragraph(paragraph, 0, 3, 1.2)
        run = paragraph.add_run(f"{number}  {item}")
        set_run_font(run, 9.7, INK, number == 4)

    note = left.add_paragraph()
    format_paragraph(note, 5, 0, 1.25)
    set_paragraph_border_and_shading(note, PALE_GOLD, GOLD)
    run = note.add_run(
        "締切前・シフト確定前なら同じ画面から修正できます。"
        "締切後、シフト確定後、突発休は先に所属管理者へ連絡してください。"
    )
    set_run_font(run, 9.2, INK, True)

    add_docx_picture(right, HOME_SCREENSHOT, 64)
    caption = right.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(caption, 0, 0, 1.1)
    caption_run = caption.add_run("① ホーム上部の希望回収を開く")
    set_run_font(caption_run, 9.1, BLUE, True)

    add_page_break(doc)
    add_heading(doc, "スマホでの取得方法", 1)
    add_text(
        doc,
        "画面の番号どおりに進めます。表示される期間・日数はスタッフごとに異なります。",
        size=9.5,
        color=MUTED,
        after=5,
    )

    screen_table = doc.add_table(rows=1, cols=2)
    set_table_geometry(screen_table, [5040, 5040])
    request_cell, balance_cell = screen_table.rows[0].cells
    for cell in (request_cell, balance_cell):
        set_cell_border(cell)
        set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
        set_cell_shading(cell, "F8FAFC")
    add_docx_picture(request_cell, REQUEST_SCREENSHOT, 75)
    request_caption = request_cell.add_paragraph()
    request_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(request_caption, 2, 0, 1.15)
    request_run = request_caption.add_run("② 種類を選ぶ → ③ 日付を押す → ④ 保存")
    set_run_font(request_run, 8.9, BLUE, True)

    add_docx_picture(balance_cell, BALANCE_SCREENSHOT, 75)
    balance_caption = balance_cell.add_paragraph()
    balance_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(balance_caption, 2, 0, 1.15)
    balance_run = balance_caption.add_run("⑤ 「有給・欠勤」で残日数を確認")
    set_run_font(balance_run, 8.9, BLUE, True)

    add_heading(doc, "覚えておくこと", 1)
    compact_notes = [
        ("半休", "実際に勤務する時間は、通常どおり出勤・退勤を打刻します。"),
        ("突発休", "まず所属管理者へ連絡します。管理者がTSGへ登録します。"),
        ("打刻がない日", "ホームの確認通知から、実際の理由を回答します。"),
        ("表示が違う", "残日数・入社日・勤務形態に疑問がある場合は、申請前に所属管理者へ連絡します。"),
    ]
    note_table = doc.add_table(rows=0, cols=2)
    set_table_geometry(note_table, [1900, 8180])
    for label, detail in compact_notes:
        cells = note_table.add_row().cells
        set_cell_shading(cells[0], PALE_BLUE)
        for index, value in enumerate((label, detail)):
            set_cell_border(cells[index])
            paragraph = cells[index].paragraphs[0]
            format_paragraph(paragraph, 0, 0, 1.12)
            run = paragraph.add_run(value)
            set_run_font(run, 8.9, INK, index == 0)
    set_table_geometry(note_table, [1900, 8180])

    add_text(
        doc,
        "口頭・電話・DMで連絡しただけでは、TSG上の有給申請は完了しません。"
        "緊急の欠勤連絡は、これまでどおり所属管理者へ直接連絡してください。",
        size=8.8,
        color=MUTED,
        after=3,
    )
    add_text(doc, "株式会社テクニカルスタッフ", size=9.4, color=NAVY, bold=True, after=0)

    properties = doc.core_properties
    properties.title = "TSG有給管理システム 全スタッフ向け運用案内"
    properties.subject = "2026年8月1日運用開始"
    properties.author = "株式会社テクニカルスタッフ"
    properties.keywords = "TSG, 有給, 年次有給休暇, 運用案内"

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


def add_title_block_compact(doc):
    add_text(doc, "全スタッフ向け運用案内", size=9.5, color=TEAL, bold=True, after=2)
    title = doc.add_paragraph()
    format_paragraph(title, 0, 3, 1.0)
    run = title.add_run("TSG有給管理システム")
    set_run_font(run, 23, NAVY, True)
    subtitle = doc.add_paragraph()
    format_paragraph(subtitle, 0, 8, 1.1)
    run = subtitle.add_run("2026年8月1日 運用開始")
    set_run_font(run, 13, BLUE, True)
    add_text(
        doc,
        "有給の確認と申請をTSGで行います。操作は「種類を選ぶ・日付を押す・保存」の3つです。",
        size=10.5,
        after=8,
    )


def pdf_phone_image(image_path: Path, width_mm: float):
    image = RLImage(str(image_path))
    image.drawWidth = width_mm * mm
    image.drawHeight = width_mm * (844 / 390) * mm
    return image


def build_compact_pdf():
    register_pdf_fonts()
    styles = pdf_styles()
    document = BaseDocTemplate(
        str(PDF_OUTPUT),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=17 * mm,
        bottomMargin=14 * mm,
        title="TSG有給管理システム 全スタッフ向け運用案内",
        author="株式会社テクニカルスタッフ",
        subject="2026年8月1日運用開始",
    )
    frame = Frame(
        document.leftMargin,
        document.bottomMargin,
        document.width,
        document.height,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )
    document.addPageTemplates([PageTemplate(id="staff-guide", frames=[frame], onPage=pdf_header_footer)])

    change_items = [
        "残日数・付与予定・取得履歴をTSGで確認できます。",
        "予定して取る有給は、シフト希望回収と一緒に提出します。",
        "有給は全休1日、半休0.5日の2種類です。",
        "8月1日の付与分は、当日になると残日数へ加算されます。",
    ]
    steps = [
        "ホーム上部の「シフト希望回収」を開く",
        "「有給（全休）」または「有給（半休）」を選ぶ",
        "希望日を押す",
        "「保存」を押して提出する",
    ]
    left_story = [
        Paragraph("8月1日から変わること", styles["h1"]),
        pdf_bullets(styles, change_items),
        Paragraph("有給を取るとき", styles["h1"]),
        pdf_steps(styles, steps),
        Paragraph(
            "<b>締切後・突発休</b><br/>"
            "締切前・シフト確定前なら同じ画面から修正できます。"
            "締切後、シフト確定後、突発休は先に所属管理者へ連絡してください。",
            styles["small"],
        ),
    ]
    home_panel = [
        pdf_phone_image(HOME_SCREENSHOT, 63),
        Paragraph("① ホーム上部の希望回収を開く", styles["table_bold"]),
    ]
    intro_table = Table(
        [[left_story, home_panel]],
        colWidths=[108 * mm, 70 * mm],
        hAlign="LEFT",
    )
    intro_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), pdf_color("F8FAFC")),
        ("BACKGROUND", (1, 0), (1, 0), pdf_color(PALE_BLUE)),
        ("BOX", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("INNERGRID", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (1, 0), (1, 0), "CENTER"),
    ]))

    request_panel = [
        pdf_phone_image(REQUEST_SCREENSHOT, 76),
        Paragraph("② 種類を選ぶ → ③ 日付を押す → ④ 保存", styles["table_bold"]),
    ]
    balance_panel = [
        pdf_phone_image(BALANCE_SCREENSHOT, 76),
        Paragraph("⑤ 「有給・欠勤」で残日数を確認", styles["table_bold"]),
    ]
    screen_table = Table(
        [[request_panel, balance_panel]],
        colWidths=[89 * mm, 89 * mm],
        hAlign="LEFT",
    )
    screen_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), pdf_color("F8FAFC")),
        ("BOX", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("INNERGRID", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))

    note_rows = [
        [Paragraph("半休", styles["table_bold"]), Paragraph("実際に勤務する時間は、通常どおり出勤・退勤を打刻します。", styles["table"])],
        [Paragraph("突発休", styles["table_bold"]), Paragraph("まず所属管理者へ連絡します。管理者がTSGへ登録します。", styles["table"])],
        [Paragraph("打刻がない日", styles["table_bold"]), Paragraph("ホームの確認通知から、実際の理由を回答します。", styles["table"])],
        [Paragraph("表示が違う", styles["table_bold"]), Paragraph("申請前に所属管理者へ連絡してください。", styles["table"])],
    ]
    note_table = Table(note_rows, colWidths=[32 * mm, 146 * mm], hAlign="LEFT")
    note_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), pdf_color(PALE_BLUE)),
        ("GRID", (0, 0), (-1, -1), 0.45, pdf_color(BORDER)),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))

    story = [
        Spacer(1, 4 * mm),
        Paragraph("全スタッフ向け運用案内", styles["kicker"]),
        Paragraph("TSG有給管理システム", styles["title"]),
        Paragraph("2026年8月1日 運用開始", styles["subtitle"]),
        Paragraph(
            "有給の確認と申請をTSGで行います。操作は「種類を選ぶ・日付を押す・保存」の3つです。",
            styles["body"],
        ),
        Spacer(1, 3 * mm),
        intro_table,
        PageBreak(),
        Paragraph("スマホでの取得方法", styles["h1"]),
        Paragraph(
            "画面の番号どおりに進めます。表示される期間・日数はスタッフごとに異なります。",
            styles["small"],
        ),
        Spacer(1, 2 * mm),
        screen_table,
        Paragraph("覚えておくこと", styles["h1"]),
        note_table,
        Spacer(1, 2 * mm),
        Paragraph(
            "口頭・電話・DMで連絡しただけでは、TSG上の有給申請は完了しません。"
            "緊急の欠勤連絡は、これまでどおり所属管理者へ直接連絡してください。",
            styles["small"],
        ),
        Spacer(1, 2 * mm),
        Paragraph('<font name="BIZUDGothic-Bold" color="#17324D">株式会社テクニカルスタッフ</font>', styles["body"]),
    ]
    PDF_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.build(story)
    print(PDF_OUTPUT)


if __name__ == "__main__":
    build_compact()
    build_compact_pdf()
